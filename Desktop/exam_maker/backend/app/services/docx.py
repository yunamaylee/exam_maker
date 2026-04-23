from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import json


# 시험지 docx 생성 함수
def create_exam_docx(exam_content: str) -> bytes:
    exam = json.loads(exam_content) if isinstance(exam_content, str) else exam_content

    doc = Document()

    # 페이지 여백 설정
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 헤더 (학교명, 시험 정보)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{exam.get('school', '')} 영어 시험지")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    # 문제 목록
    for question in exam.get('questions', []):
        # 문제 번호 및 유형
        header = doc.add_paragraph()
        run = header.add_run(
            f"{question['number']}. [{question['type']}] ({question['score']}점)"
        )
        run.bold = True
        run.font.size = Pt(11)

        # 지문
        passage = doc.add_paragraph()
        passage.add_run(question.get('passage', ''))
        passage.paragraph_format.left_indent = Cm(0.5)

        # 문제
        q_para = doc.add_paragraph()
        run = q_para.add_run(question.get('question', ''))
        run.bold = True

        # 선택지 (객관식)
        if question.get('choices'):
            for choice in question['choices']:
                choice_para = doc.add_paragraph()
                choice_para.add_run(choice)
                choice_para.paragraph_format.left_indent = Cm(0.5)

        # 조건 (서술형)
        if question.get('condition'):
            condition_para = doc.add_paragraph()
            run = condition_para.add_run('[조건]')
            run.bold = True
            for cond in question['condition']:
                c_para = doc.add_paragraph()
                c_para.add_run(f'• {cond}')
                c_para.paragraph_format.left_indent = Cm(0.5)

        doc.add_paragraph()

    # 정답 구분선
    doc.add_paragraph('─' * 50)
    answer_title = doc.add_paragraph()
    run = answer_title.add_run('【정답】')
    run.bold = True
    run.font.size = Pt(12)

    # 정답 목록
    for question in exam.get('questions', []):
        answer_para = doc.add_paragraph()
        answer_para.add_run(
            f"{question['number']}번: {question.get('answer', '')}"
        )

    # bytes로 변환
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()