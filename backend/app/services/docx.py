from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import json


def set_two_columns(section):
    """섹션을 2단으로 설정"""
    sectPr = section._sectPr
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '708')  # 단 사이 간격
    sectPr.append(cols)


def add_horizontal_line(doc):
    """구분선 추가"""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


# 시험지 docx 생성 함수
def create_exam_docx(exam_content: str) -> bytes:
    exam = json.loads(exam_content) if isinstance(exam_content, str) else exam_content

    doc = Document()

    # 기본 폰트 설정
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)

    # 페이지 여백 설정
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # 헤더 - 학교명
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{exam.get('school', '')} 영어 시험지")
    run.bold = True
    run.font.size = Pt(14)

    # 부제목 (총점)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"총점: {exam.get('total_score', 100)}점")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    add_horizontal_line(doc)
    doc.add_paragraph()

    # 2단 설정
    set_two_columns(doc.sections[0])

    # 문제 목록
    questions = exam.get('questions', [])
    mc_questions = [q for q in questions if q['type'] == '객관식']
    sub_questions = [q for q in questions if q['type'] == '서술형']

    # 객관식 문제
    for question in mc_questions:
        # 문제 번호
        header = doc.add_paragraph()
        run = header.add_run(
            f"{question['number']}. ({question['score']}점)"
        )
        run.bold = True
        run.font.size = Pt(10)

        # 문제 지시문
        q_para = doc.add_paragraph()
        run = q_para.add_run(question.get('question', ''))
        run.bold = True
        run.font.size = Pt(9)

        # 지문
        if question.get('passage'):
            passage = doc.add_paragraph()
            run = passage.add_run(question.get('passage', ''))
            run.font.size = Pt(9)
            passage.paragraph_format.left_indent = Cm(0.3)
            passage.paragraph_format.space_before = Pt(2)
            passage.paragraph_format.space_after = Pt(2)

        # 선택지
        if question.get('choices'):
            for choice in question['choices']:
                choice_para = doc.add_paragraph()
                run = choice_para.add_run(choice)
                run.font.size = Pt(9)
                choice_para.paragraph_format.left_indent = Cm(0.3)

        doc.add_paragraph()

    # 서술형 구분선
    add_horizontal_line(doc)
    sub_title = doc.add_paragraph()
    run = sub_title.add_run('【서술형】')
    run.bold = True
    run.font.size = Pt(11)
    doc.add_paragraph()

    # 서술형 문제
    for question in sub_questions:
        # 문제 번호
        header = doc.add_paragraph()
        run = header.add_run(
            f"서답형 {question['number'] - len(mc_questions)}. ({question['score']}점)"
        )
        run.bold = True
        run.font.size = Pt(10)

        # 문제 지시문
        q_para = doc.add_paragraph()
        run = q_para.add_run(question.get('question', ''))
        run.bold = True
        run.font.size = Pt(9)

        # 지문
        if question.get('passage'):
            passage = doc.add_paragraph()
            run = passage.add_run(question.get('passage', ''))
            run.font.size = Pt(9)
            passage.paragraph_format.left_indent = Cm(0.3)

        # 조건
        if question.get('condition'):
            condition_title = doc.add_paragraph()
            run = condition_title.add_run('[조건]')
            run.bold = True
            run.font.size = Pt(9)
            for cond in question['condition']:
                c_para = doc.add_paragraph()
                run = c_para.add_run(f'• {cond}')
                run.font.size = Pt(9)
                c_para.paragraph_format.left_indent = Cm(0.3)

        # 답안 작성란
        answer_label = doc.add_paragraph()
        run = answer_label.add_run('[답]')
        run.bold = True
        run.font.size = Pt(9)

        # 빈 줄 3개 (답안 작성 공간)
        for _ in range(3):
            blank = doc.add_paragraph()
            blank.paragraph_format.space_before = Pt(8)
            add_horizontal_line(blank.part.document if hasattr(blank.part, 'document') else doc)

        doc.add_paragraph()

    # 정답 페이지
    doc.add_page_break()
    answer_title = doc.add_paragraph()
    run = answer_title.add_run('【정답 및 해설】')
    run.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph()

    for question in questions:
        answer_para = doc.add_paragraph()
        run = answer_para.add_run(
            f"{question['number']}번: {question.get('answer', '')}"
        )
        run.font.size = Pt(10)

    # bytes로 변환
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()